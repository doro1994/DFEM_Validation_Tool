# -*- coding: utf-8 -*-
"""
Created on Mon Jun 20 11:05:03 2022

@author: DONIK94
"""

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from dfem_validation.f06_processing import analyse_f06
import re


def write_output_from_gpwg(document, output_from_gpwg):
    p = document.add_paragraph("The mass output is depicted in table below")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in output_from_gpwg["lines"]:
        run = p.add_run(line.strip() + "\n")
        run.font.name = "Courier"
        run.font.size = Pt(8)
    p.add_run(f"Output from Grid Point Weight Generator \n").bold = True
    return document


def write_OLOAD_resultant(document, OLOAD_RESULTANT, checkname):
    p = document.add_paragraph("OLOAD Resultant output is shown in the following table")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for header_line in OLOAD_RESULTANT["header"]:
        run = p.add_run(header_line.rstrip() + "\n")
        run.font.name = "Courier"
        run.font.size = Pt(8)
    for subcase in OLOAD_RESULTANT["lines"]:
        for line in OLOAD_RESULTANT["lines"][subcase]:
            run = p.add_run(line.rstrip() + "\n")
            run.font.name = "Courier"
            run.font.size = Pt(8)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"OLOAD Resultant for {checkname} Check  \n").bold = True
    return document


def write_SPCFORCE_resultant(document, SPCFORCE_RESULTANT, checkname):
    p = document.add_paragraph(
        "SPCFORCE Resultant output is shown in the following table"
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for header_line in SPCFORCE_RESULTANT["header"]:
        run = p.add_run(header_line.rstrip() + "\n")
        run.font.name = "Courier"
        run.font.size = Pt(8)
    for subcase in SPCFORCE_RESULTANT["lines"]:
        for line in SPCFORCE_RESULTANT["lines"][subcase]:
            run = p.add_run(line.rstrip() + "\n")
            run.font.name = "Courier"
            run.font.size = Pt(8)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"SPCFORCE Resultant for {checkname} Check  \n").bold = True
    return document


def add_picture(document, picture_filepath, picture_name):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(picture_filepath, width=Inches(6.7))
    p.add_run(f"{picture_name}\n").bold = True
    return document


def get_picture_filenames(
    check_name, screenshots_folder="./DFEM_Validation/screenshots"
):
    picture_filenames = []
    exatrcted_numbers = []
    for path in Path(screenshots_folder).iterdir():
        if path.is_file() and check_name.lower() in path.stem:
            picture_filenames.append(path)
            exatrcted_numbers.append(extract_number(str(path)))

    picture_filenames.sort(key=extract_number)
    return picture_filenames


def extract_number(string):
    number = re.findall(r"\d+", str(string))[0]
    return int(number)


def write_external_work(document, EXTERNAL_WORK, checkname):
    p = document.add_paragraph(
        "SPCFORCE Resultant output is shown in the following table"
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for line in EXTERNAL_WORK["lines"]:
        run = p.add_run(line.rstrip() + "\n")
        run.font.name = "Courier"
        run.font.size = Pt(8)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"Matrix-to-Factor-Diagonal, Epsilon, External Work for {checkname} Check  \n"
    ).bold = True

    return document


def write_eigenvalues(
    document, f06_result_dict_eig, checkname, rigid_only=False, flexible_only=False
):

    p = document.add_paragraph(
        "A free-free modal analysis is performed to verify that there are six Rigid Body Modes (RBM). The LANCZOS method is used."
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    freq_num = 0
    freq_max = 0
    for mode in f06_result_dict_eig["frequencies"]:
        freq = float(f06_result_dict_eig["frequencies"][mode])
        freq_num += 1
        if freq > 0.01:
            print(f"Frequency for mode {mode} is greater than 0.01")
            print(f"F = {freq} Hz")
            return None
        if freq > float(freq_max):
            freq_max = f06_result_dict_eig["frequencies"][mode]
        if freq_num >= 6:
            break
    p.add_run(
        f"The highest RBM frequency of {freq_max} Hz is below required value of 0.01 Hz \n"
    )

    p = document.add_paragraph()

    run = p.add_run(
        f"                                              R E A L   E I G E N V A L U E S \n"
        f"MODE  EXTRACTION  EIGENVALUE        RADIANS          CYCLES        GENERALIZED      GENERALIZED \n"
        f" NO.    ORDER                                                          MASS          STIFFNESS \n"
    )

    run.font.name = "Courier"
    run.font.size = Pt(8)
    for mode in f06_result_dict_eig["lines"]:
        if rigid_only:
            if mode > 6:
                break
        if flexible_only:
            if mode <= 6:
                continue

        line = f06_result_dict_eig["lines"][mode].replace("        ", "     ")
        if line.startswith("       "):
            line = "    " + line[7:]

        run = p.add_run(line.rstrip() + "\n")
        run.font.name = "Courier"
        run.font.size = Pt(8)

    if rigid_only and not flexible_only:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Summary of Eigenvalues of 6 RBM\n").bold = True
    elif flexible_only and not rigid_only:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Summary of 15 flexible Eigenvalues\n").bold = True
    elif rigid_only and flexible_only:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Summary of Eigenvalues \n").bold = True

    return document


def get_check_name(f06_filename):
    if "gravity" in f06_filename.lower():
        return "Gravity"
    elif "unit" in f06_filename.lower():
        return "Unit Displacement"
    elif "dynamic" in f06_filename.lower():
        return "Dynamic"
    elif "thermoelastic" in f06_filename.lower():
        return "Thermoelastic"


def create_document(document_heading="DFEM Validation Preliminary Report"):
    document = Document()
    document.add_heading(document_heading, 0)

    # changing the page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.59)
    return document


def write_gravity_check_results(document, f06_gravitycheck_dict, screenshots_folder):
    document.add_heading("Gravity Check Results", level=1)
    output_document = write_output_from_gpwg(
        document, f06_gravitycheck_dict["output_from_gpwg"]
    )
    output_document = write_OLOAD_resultant(
        document, f06_gravitycheck_dict["OLOAD    RESULTANT"], checkname="Gravity"
    )
    output_document = write_SPCFORCE_resultant(
        document, f06_gravitycheck_dict["SPCFORCE RESULTANT"], checkname="Gravity"
    )

    if f06_gravitycheck_dict["matrix_to_factor_diag_ratio"]["value"] < 1e7:
        output_document.add_paragraph(
            f"As required, the Matrix-to-Factor-Diagonal is below 1.0E07. Further, Epsilon and External Work are shown in table below"
        )
    else:
        output_document.add_paragraph(
            f"ERROR! Matrix-to-Factor-Diagonal is above 1.0E07. Epsilon and External Work are shown in table below"
        )

    output_document = write_external_work(
        output_document, f06_gravitycheck_dict["EXTERNAL WORK"], checkname="Gravity"
    )

    output_document.add_paragraph(
        f"Figures below show displacements for the three gravity load cases which are reasonably small. "
        f"Thus, it can be deduced that all the parts are tightly connected with each other."
    )

    # Get screenshots
    picture_filenames = get_picture_filenames(
        check_name="Gravity", screenshots_folder=screenshots_folder
    )
    load_cases = ["X", "Y", "Z"]
    for pic_id in range(len(picture_filenames)):

        output_document = add_picture(
            document=output_document,
            picture_filepath=str(picture_filenames[pic_id]),
            picture_name=f"Gravity Load Case 1g {load_cases[pic_id]}, Displacement",
        )

    output_document.add_page_break()
    return output_document


def write_unit_check_results(document, f06_unitcheck_dict, screenshots_folder):
    document.add_heading("Unit Displacement Check Results", level=1)
    output_document = write_OLOAD_resultant(
        document, f06_unitcheck_dict["OLOAD    RESULTANT"], checkname="Unit"
    )
    output_document = write_SPCFORCE_resultant(
        document, f06_unitcheck_dict["SPCFORCE RESULTANT"], checkname="Unit"
    )

    output_document.add_paragraph(
        f"Figures below prove that the structure behavior is stress-free, as required."
    )

    # Get screenshots
    picture_filenames = get_picture_filenames(
        check_name="Unit", screenshots_folder=screenshots_folder
    )
    load_cases = ["tx", "ty", "tz", "rx", "ry", "rz"]
    for pic_id in range(len(picture_filenames)):

        output_document = add_picture(
            document=output_document,
            picture_filepath=str(picture_filenames[pic_id]),
            picture_name=f"Unit Dispalcement Load Case {load_cases[pic_id]}, v. Mises Stress",
        )

    output_document.add_page_break()
    return output_document


def write_dynamic_check_results(document, f06_dynamic_dict, screenshots_folder):
    document.add_heading("Dynamic Check Results", level=1)
    output_document = write_eigenvalues(
        document,
        f06_dynamic_dict["R E A L   E I G E N V A L U E S"],
        checkname="Dynamic",
        rigid_only=True,
        flexible_only=False,
    )

    output_document.add_paragraph(f"Six Rigid Body Modes are shown below")

    # Get screenshots
    picture_filenames = get_picture_filenames(
        check_name="Dynamic", screenshots_folder=screenshots_folder
    )
    for pic_id in range(6):

        output_document = add_picture(
            document=output_document,
            picture_filepath=str(picture_filenames[pic_id]),
            picture_name=f"Rigid Body Mode {pic_id + 1}",
        )

    if (
        float(f06_dynamic_dict["R E A L   E I G E N V A L U E S"]["frequencies"][7])
        / float(f06_dynamic_dict["R E A L   E I G E N V A L U E S"]["frequencies"][6])
    ) > 10000:
        output_document.add_paragraph(
            f"The ratio of the 1st elastic mode and the last RBM is larger than 1.0E+4 as required."
        )
    else:
        output_document.add_paragraph(
            f"WARNING!!! The ratio of the 1st elastic mode and the last RBM is below 1.0E+4."
        )

    output_document.add_paragraph(
        f"The summary for the first 15 flexible modes is shown in the follwoing table:"
    )

    output_document = write_eigenvalues(
        document,
        f06_dynamic_dict["R E A L   E I G E N V A L U E S"],
        checkname="Dynamic",
        rigid_only=False,
        flexible_only=True,
    )

    output_document.add_paragraph(
        f"The first 15 flexible modes are shown in the figures below"
    )
    for pic_id in range(6, 21):

        output_document = add_picture(
            document=output_document,
            picture_filepath=str(picture_filenames[pic_id]),
            picture_name=f"Flexible Mode {pic_id + 1 - 6}",
        )
    output_document.add_page_break()
    return output_document


def write_thermoelastic_check_results(
    document, f06_thermoelastic_dict, screenshots_folder
):
    document.add_heading("Thermoelastic Check Results", level=1)

    document.add_paragraph(f"The displacements are shown in the following Figure")

    # Get screenshots
    picture_filenames = get_picture_filenames(
        check_name="Thermoelastic", screenshots_folder=screenshots_folder
    )

    output_document = add_picture(
        document=document,
        picture_filepath=str(picture_filenames[0]),
        picture_name=f"Elasto-Thermic Check, Displacements",
    )

    output_document = add_picture(
        document=document,
        picture_filepath=str(picture_filenames[1]),
        picture_name=f"Elasto-Thermic Check, v.Mises",
    )

    output_document.add_paragraph(f"OLOAD and SPCFORCE resultants are the following:")
    output_document = write_OLOAD_resultant(
        document,
        f06_thermoelastic_dict["OLOAD    RESULTANT"],
        checkname="Thermoelastic",
    )
    output_document = write_SPCFORCE_resultant(
        document,
        f06_thermoelastic_dict["SPCFORCE RESULTANT"],
        checkname="Thermoelastic",
    )

    output_document.add_page_break()
    return output_document


def create_report(
    f06_gravitycheck_dict,
    f06_unitcheck_dict,
    f06_dynamic_dict,
    f06_thermoelastic_dict,
    filename="DFEM_Validation_Preliminary_report.docx",
    screenshots_folder="./DFEM_Validation/screenshots",
):

    doc = create_document(document_heading="DFEM Validation Preliminary Report")
    doc = write_gravity_check_results(
        doc, f06_gravitycheck_dict, screenshots_folder=screenshots_folder
    )
    doc = write_unit_check_results(
        doc, f06_unitcheck_dict, screenshots_folder=screenshots_folder
    )
    doc = write_dynamic_check_results(
        doc, f06_dynamic_dict, screenshots_folder=screenshots_folder
    )
    doc = write_thermoelastic_check_results(
        doc, f06_thermoelastic_dict, screenshots_folder=screenshots_folder
    )

    doc.save(filename)
